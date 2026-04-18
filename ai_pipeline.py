import hashlib
import os
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

import pdfplumber

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_community.vectorstores import Chroma
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

try:
    from langchain_community.chat_models import ChatOllama
    from langchain_community.embeddings import OllamaEmbeddings
except Exception:  # pragma: no cover
    ChatOllama = None  # type: ignore
    OllamaEmbeddings = None  # type: ignore

from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer

from database import Database


@dataclass
class RetrievedChunk:
    text: str
    metadata: Dict[str, Any]
    score: Optional[float] = None


def sha256_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _utcnow_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def file_to_documents(file_path: str) -> Tuple[List[Document], str, Optional[str]]:
    file_name = os.path.basename(file_path)
    ext = os.path.splitext(file_name)[1].lower()

    docs: List[Document] = []
    if ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source_file": file_name,
                                "page": i + 1,
                                "doc_type": "pdf",
                            },
                        )
                    )
    elif ext == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
        if full_text.strip():
            docs.append(
                Document(
                    page_content=full_text,
                    metadata={
                        "source_file": file_name,
                        "page": None,
                        "doc_type": "docx",
                    },
                )
            )
    else:
        # Assume .txt (or unknown) falls back to plain text
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source_file": file_name,
                        "page": None,
                        "doc_type": "text",
                    },
                )
            )

    if not docs:
        # Return empty docs; caller can decide error handling.
        return [], file_name, ext

    return docs, file_name, ext


def chunk_documents(
    docs: List[Document],
    chunk_size: int = 1200,
    chunk_overlap: int = 200,
) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    return splitter.split_documents(docs)




def get_vectorstore(
    user_id: int,
    embeddings,
    persist_base_dir: str,
):
    persist_dir = os.path.join(persist_base_dir, str(user_id))
    os.makedirs(persist_dir, exist_ok=True)

    return Chroma(
        collection_name="kb_collection",
        embedding_function=embeddings,
        persist_directory=persist_dir,   # ✅ THIS IS THE FIX
    )
#def get_vectorstore(
#    user_id: int,
#    embeddings,
#    persist_base_dir: str,
#):
#    persist_dir = os.path.join(persist_base_dir, str(user_id))
#    os.makedirs(persist_dir, exist_ok=True)
#    # Create an explicit persistent client so we don't accidentally create an
#    # in-memory ("ephemeral") client with different settings.
#    
#
#    return Chroma(
#        # Chroma collection name must be 3-512 chars and start/end with [a-zA-Z0-9].
#        collection_name="kb_collection",
#        embedding_function=embeddings,
#        client=client,
#    )


def get_embeddings(
    llm_provider: str,
    embeddings_provider: Optional[str] = None,
    ollama_base_url: Optional[str] = None,
):
    provider = embeddings_provider or llm_provider

    if provider == "openai":
        # Small embeddings are faster; for demo quality this is usually sufficient.
        return OpenAIEmbeddings(model="text-embedding-3-small")

    if provider == "ollama":
        if OllamaEmbeddings is None:
            raise RuntimeError("Ollama embeddings not available (install langchain-community).")
        base_url = ollama_base_url or "http://localhost:11434"
        return OllamaEmbeddings(base_url=base_url, model="nomic-embed-text")

    # Default to OpenAI embeddings if no provider is set.
    return OpenAIEmbeddings(model="text-embedding-3-small")


def get_chat_model(
    llm_provider: str,
    openai_api_key: Optional[str] = None,
    ollama_base_url: Optional[str] = None,
):
    if llm_provider == "ollama":
        if ChatOllama is None:
            raise RuntimeError("ChatOllama not available (install langchain-community).")
        base_url = ollama_base_url or "http://localhost:11434"
        return ChatOllama(base_url=base_url, model="llama3", temperature=0.2)

    # openai default
    return ChatOpenAI(
        temperature=0.2,
        model="gpt-4o-mini",
        api_key=openai_api_key,
    )


def ingest_document(
    db: Database,
    user_id: int,
    file_path: str,
    persist_base_dir: str,
    embeddings,
    llm_provider: str,
    force_reingest: bool = False,
) -> Dict[str, Any]:
    if embeddings is None:
        # Let env control which embeddings backend to use.
        embeddings = get_embeddings(
            llm_provider=llm_provider,
            embeddings_provider=os.getenv("EMBEDDINGS_PROVIDER") or None,
            ollama_base_url=os.getenv("OLLAMA_BASE_URL") or None,
        )
    doc_hash = sha256_file(file_path)
    already = db.doc_exists(user_id=user_id, doc_hash=doc_hash)
    if already and not force_reingest:
        return {"status": "cached", "doc_hash": doc_hash}

    raw_docs, file_name, ext = file_to_documents(file_path)
    if not raw_docs:
        raise ValueError("No extractable text found in the uploaded document.")

    chunks = chunk_documents(raw_docs)
    vectorstore = get_vectorstore(user_id, embeddings, persist_base_dir)
    ids = [f"{doc_hash}_{i}" for i in range(len(chunks))]

    # Add the chunks with doc_hash in their metadata for traceability.
    for i, d in enumerate(chunks):
        d.metadata = dict(d.metadata or {})
        d.metadata.update(
            {
                "doc_hash": doc_hash,
                "chunk_index": i,
                "ingested_at": _utcnow_iso(),
            }
        )

    vectorstore.add_documents(documents=chunks, ids=ids)
    vectorstore.persist()

    # Persist metadata in SQL for caching and deletion.
    from mimetypes import guess_type

    mime, _ = guess_type(file_path)
    if not already:
        db.add_document(
            user_id=user_id,
            doc_hash=doc_hash,
            file_name=file_name,
            file_mime=mime,
            chunk_count=len(chunks),
        )
        db.log_upload(user_id=user_id, doc_hash=doc_hash, file_name=file_name)

    return {"status": "ingested", "doc_hash": doc_hash, "chunk_count": len(chunks), "ext": ext}


def retrieve_chunks(
    vectorstore,
    query: str,
    k: int = 5,
) -> List[RetrievedChunk]:
    # similarity_search_with_relevance_scores is more robust, but score meaning may vary.
    results = vectorstore.similarity_search_with_relevance_scores(query, k=k)
    out: List[RetrievedChunk] = []
    for doc, score in results:
        out.append(RetrievedChunk(text=doc.page_content, metadata=dict(doc.metadata or {}), score=float(score)))
    return out


def build_rag_prompt(query: str, chunks: List[RetrievedChunk]) -> str:
    context_blocks: List[str] = []
    for idx, ch in enumerate(chunks, start=1):
        src_file = ch.metadata.get("source_file", "unknown")
        page = ch.metadata.get("page", None)
        page_str = f"page {page}" if page else "page unknown"
        context_blocks.append(
            f"[Context {idx}] Source={src_file} ({page_str})\n{ch.text}\n"
        )

    context = "\n".join(context_blocks) if context_blocks else "No relevant context found."

    return f"""You are a helpful AI assistant. Answer the user's question using ONLY the provided context.
If the context is insufficient, say you don't know and suggest what to upload or how to ask.

User question:
{query}

Provided context:
{context}

Rules:
- Be concise but complete.
- Cite sources by writing "Sources:" followed by a comma-separated list like "filename (page X)" for the contexts you used.
"""


def analyze_sentiment(text: str) -> Dict[str, Any]:
    try:
        analyzer = SentimentIntensityAnalyzer()
        scores = analyzer.polarity_scores(text)
        compound = float(scores.get("compound", 0.0))
        if compound >= 0.05:
            label = "positive"
        elif compound <= -0.05:
            label = "negative"
        else:
            label = "neutral"
        return {"label": label, "compound": compound, "scores": scores}
    except Exception:
        # Keep it resilient for demo environments.
        return {"label": "neutral", "compound": 0.0, "scores": {}}


def summarize_text(
    text: str,
    llm_provider: str,
    openai_api_key: Optional[str] = None,
    ollama_base_url: Optional[str] = None,
    max_chars: int = 9000,
) -> str:
    """
    Summarize arbitrary text using the selected LLM provider.
    """
    truncated = text[:max_chars]
    chat_model = get_chat_model(
        llm_provider=llm_provider,
        openai_api_key=openai_api_key,
        ollama_base_url=ollama_base_url,
    )
    system = (
        "You are an expert study assistant. Provide a concise, accurate summary. "
        "Use bullet points and keep it grounded in the source text."
    )
    prompt = f"{system}\n\nText:\n{truncated}\n\nSummary:"
    resp = chat_model.invoke(prompt)
    return getattr(resp, "content", str(resp)).strip()


def summarize_chat(
    messages: List[Dict[str, Any]],
    llm_provider: str,
    openai_api_key: Optional[str] = None,
    ollama_base_url: Optional[str] = None,
) -> str:
    transcript_parts: List[str] = []
    for m in messages:
        role = (m.get("role") or "unknown").lower()
        content = m.get("content") or ""
        transcript_parts.append(f"{role.upper()}: {content}")
    transcript = "\n\n".join(transcript_parts)
    return summarize_text(
        text=transcript,
        llm_provider=llm_provider,
        openai_api_key=openai_api_key,
        ollama_base_url=ollama_base_url,
    )


def preview_document_text(file_path: str, max_chars: int = 2000) -> str:
    """
    Lightweight preview for "view files" UI.
    For PDFs we show the first non-empty page(s).
    """
    file_name = os.path.basename(file_path)
    ext = os.path.splitext(file_name)[1].lower()
    try:
        if ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                parts: List[str] = []
                for i, page in enumerate(pdf.pages):
                    txt = page.extract_text() or ""
                    if txt.strip():
                        parts.append(txt.strip())
                    if sum(len(p) for p in parts) >= max_chars or i >= 4:
                        break
                return "\n\n".join(parts)[:max_chars]

        if ext == ".docx":
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            parts: List[str] = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
                if sum(len(x) for x in parts) >= max_chars:
                    break
            return "\n\n".join(parts)[:max_chars]

        # TXT fallback
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()[:max_chars]
    except Exception as e:
        return f"Preview not available: {e}"


def answer_with_rag(
    db: Database,
    user_id: int,
    query: str,
    persist_base_dir: str,
    embeddings,
    llm_provider: str,
    openai_api_key: Optional[str] = None,
    ollama_base_url: Optional[str] = None,
    top_k: int = 5,
) -> Dict[str, Any]:
    if embeddings is None:
        embeddings = get_embeddings(
            llm_provider=llm_provider,
            embeddings_provider=os.getenv("EMBEDDINGS_PROVIDER") or None,
            ollama_base_url=ollama_base_url,
        )
    vectorstore = get_vectorstore(user_id, embeddings, persist_base_dir)
    chunks: List[RetrievedChunk] = []
    try:
        chunks = retrieve_chunks(vectorstore, query=query, k=top_k) if vectorstore is not None else []
    except Exception:
        chunks = []
    if not chunks:
        return {
            "answer": "No documents are indexed for your account yet. Please upload documents first.",
            "sources": [],
            "retrieved_chunks": [],
            "sentiment": analyze_sentiment(query),
        }

    prompt = build_rag_prompt(query, chunks)
    chat_model = get_chat_model(
        llm_provider=llm_provider,
        openai_api_key=openai_api_key,
        ollama_base_url=ollama_base_url,
    )

    resp = chat_model.invoke(prompt)
    answer_text = getattr(resp, "content", str(resp))

    # Build a reliable source list from retrieved chunk metadata.
    sources: List[str] = []
    for ch in chunks:
        src_file = ch.metadata.get("source_file", "unknown")
        page = ch.metadata.get("page", None)
        page_str = f"page {page}" if page else "page unknown"
        sources.append(f"{src_file} ({page_str})")

    # Deduplicate while preserving order
    seen = set()
    sources_unique = []
    for s in sources:
        if s not in seen:
            sources_unique.append(s)
            seen.add(s)

    sentiment = analyze_sentiment(query)
    return {
        "answer": answer_text,
        "sources": sources_unique,
        "retrieved_chunks": [
            {"score": c.score, "metadata": c.metadata, "text_preview": c.text[:400]} for c in chunks
        ],
        "sentiment": sentiment,
    }

